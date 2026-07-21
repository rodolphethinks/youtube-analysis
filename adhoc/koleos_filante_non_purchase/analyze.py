"""
Ad-hoc Analysis: Why Koreans Don't Buy the Grand Koleos or Filante
------------------------------------------------------------------
Objective:
  - Understand main reasons Korean consumers avoid the Grand Koleos or Filante
  - Understand why they prefer competitor cars instead
  - Understand post-purchase regret (kept in a separate section)

Approach:
  1. Multiple Korean YouTube search queries per car (regionCode=KR, relevanceLanguage=ko)
  2. Fetch video transcripts via youtube-transcript-api (Korean captions preferred)
  3. Fetch up to 200 comments per video; pre-filter to >=2 likes
  4. Rank each video: video_rank = 0.5*(likes/max_likes) + 0.5*(views/max_views)
  5. Rank each comment: comment_rank = log(1 + video_rank) * comment_likes
  6. Extract categorized arguments from transcripts (parallel Gemini calls per video)
  7. Extract categorized arguments from comments (parallel Gemini calls, 50 comments/batch)
  8. Merge & re-rank similar arguments per category via Gemini semantic clustering
  9. Output: structured .docx report + intermediate CSVs
"""

import os
import re
import sys
import json
import time
import math
import logging
from pathlib import Path
from typing import Optional
from concurrent.futures import ThreadPoolExecutor, as_completed

import re
import tempfile
import pandas as pd
from dotenv import load_dotenv
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from google import genai
from google.genai import types as genai_types
from pydantic import BaseModel, Field
import yt_dlp
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

# ── Setup ──────────────────────────────────────────────────────────────────────
ROOT = Path(__file__).resolve().parent.parent.parent
sys.path.insert(0, str(ROOT))
load_dotenv(ROOT / ".env")

OUTPUT_DIR = Path(__file__).parent / "output"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

CACHE_DIR = OUTPUT_DIR / "extraction_cache"
CACHE_DIR.mkdir(parents=True, exist_ok=True)

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler(OUTPUT_DIR / "run.log", encoding="utf-8"),
    ],
)
log = logging.getLogger(__name__)

GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY", "")   # Gemini only
YOUTUBE_API_KEY = os.getenv("YOUTUBE_API_KEY", "") or GOOGLE_API_KEY  # fallback for backwards compat
if not GOOGLE_API_KEY:
    raise EnvironmentError("GOOGLE_API_KEY is not set in .env")
if not YOUTUBE_API_KEY:
    raise EnvironmentError("YOUTUBE_API_KEY is not set in .env")

# ── Config ─────────────────────────────────────────────────────────────────────
COOKIES_PATH = ROOT / "cookies.txt"   # export from browser via "Get cookies.txt LOCALLY"
GEMINI_MODEL = "gemini-3.1-flash-lite"
MAX_COMMENTS = 200
MIN_COMMENT_LIKES = 2
COMMENT_BATCH_SIZE = 50
GEMINI_WORKERS = 5        # parallel Gemini threads
MAX_RESULTS_PER_QUERY = 25
GEMINI_RETRY_ATTEMPTS = 3

# gemini-3.1-flash-lite standard/paid tier pricing (ai.google.dev/gemini-api/docs/pricing)
INPUT_PRICE_PER_1M = 0.25
OUTPUT_PRICE_PER_1M = 1.50

CAR_NAMES = {
    "koleos": "Renault Grand Koleos",
    "filante": "Renault Filante",
}

# Korean search queries per car — focused on non-purchase & competitor preference
SEARCH_QUERIES: dict[str, list[str]] = {
    "koleos": [
        "그랑 콜레오스 단점",               # Grand Koleos disadvantages
        "그랑 콜레오스 구매 안하는 이유",    # Reasons not to buy Grand Koleos
        "그랑 콜레오스 실망",               # Grand Koleos disappointment
        "그랑 콜레오스 후회",               # Grand Koleos regret
        "그랑 콜레오스 문제점",             # Grand Koleos problems
        "그랑 콜레오스 경쟁차 비교",         # Grand Koleos vs competitors
        "그랑 콜레오스 VS",                # Grand Koleos vs
        "르노 그랑 콜레오스 솔직 리뷰",      # Honest review
        "그랑 콜레오스 대신 추천",           # Recommended instead of Grand Koleos
        "그랑 콜레오스 비추",               # Do not recommend Grand Koleos
    ],
    "filante": [
        "필랑트 단점",                     # Filante disadvantages
        "필랑트 구매 안하는 이유",           # Reasons not to buy Filante
        "필랑트 후회",                     # Filante regret
        "필랑트 실망",                     # Filante disappointment
        "필랑트 문제점",                   # Filante problems
        "필랑트 VS",                       # Filante vs
        "르노 필랑트 솔직 리뷰",             # Honest Filante review
        "필랑트 대신 추천",                 # Recommended instead of Filante
        "필랑트 비추",                     # Do not recommend Filante
        "필랑트 경쟁차",                   # Filante competitors
    ],
}

CATEGORIES = ["non_purchase", "competitor", "regret"]
CATEGORY_LABELS = {
    "non_purchase": "Reasons Not to Buy",
    "competitor":   "Reasons to Prefer a Competitor",
    "regret":       "Post-Purchase Regret",
}


# ── YouTube helpers ────────────────────────────────────────────────────────────

def build_youtube(api_key: str):
    return build("youtube", "v3", developerKey=api_key)


def search_videos(yt, query: str, max_results: int = MAX_RESULTS_PER_QUERY) -> dict[str, str]:
    """
    Search YouTube with Korean region/language filters.
    Falls back to removing regionCode if no results are returned.
    Returns {url: video_id}.
    """
    base_kwargs = dict(
        part="snippet",
        maxResults=max_results,
        q=query,
        type="video",
        relevanceLanguage="ko",
    )
    for use_region in (True, False):
        kwargs = dict(base_kwargs)
        if use_region:
            kwargs["regionCode"] = "KR"
        try:
            resp = yt.search().list(**kwargs).execute()
            items = resp.get("items", [])
            if items:
                return {
                    f"https://www.youtube.com/watch?v={item['id']['videoId']}": item["id"]["videoId"]
                    for item in items
                }
        except HttpError as e:
            log.warning(f"Search error for '{query}' (region={use_region}): {e}")
    return {}


def get_video_details(yt, video_id: str) -> Optional[dict]:
    """Fetch snippet + statistics for a single video."""
    try:
        resp = yt.videos().list(
            part="snippet,statistics",
            id=video_id,
        ).execute()
        if not resp.get("items"):
            return None
        item = resp["items"][0]
        stats = item["statistics"]
        return {
            "video_id": video_id,
            "url": f"https://www.youtube.com/watch?v={video_id}",
            "title": item["snippet"]["title"],
            "channel": item["snippet"].get("channelTitle", ""),
            "published_at": item["snippet"]["publishedAt"][:10],
            "views": int(stats.get("viewCount", 0)),
            "likes": int(stats.get("likeCount", 0)),
            "comment_count": int(stats.get("commentCount", 0)),
        }
    except HttpError as e:
        log.warning(f"Video details error for {video_id}: {e}")
        return None


def fetch_comments(yt, video_id: str, max_comments: int = MAX_COMMENTS) -> list[dict]:
    """
    Fetch up to max_comments top-level comments ordered by relevance.
    Paginates across up to 2 pages (200 comments max).
    """
    comments: list[dict] = []
    seen_ids: set[str] = set()  # dedupe: relevance-ordered pagination can repeat items across pages

    def _parse_page(response: dict) -> None:
        for item in response.get("items", []):
            thread_id = item.get("id", "")
            if thread_id and thread_id in seen_ids:
                continue
            if thread_id:
                seen_ids.add(thread_id)
            s = item["snippet"]["topLevelComment"]["snippet"]
            comments.append({
                "comment_id": thread_id,
                "video_id": video_id,
                "author": s.get("authorDisplayName", ""),
                "comment": s.get("textDisplay", ""),
                "likes": int(s.get("likeCount", 0)),
                "reply_count": int(item["snippet"].get("totalReplyCount", 0)),
                "published_at": s.get("publishedAt", "")[:10],
            })

    try:
        resp = yt.commentThreads().list(
            part="snippet",
            videoId=video_id,
            maxResults=100,
            order="relevance",
            textFormat="plainText",
        ).execute()
        _parse_page(resp)

        page_token = resp.get("nextPageToken")
        if page_token and len(comments) < max_comments:
            resp2 = yt.commentThreads().list(
                part="snippet",
                videoId=video_id,
                maxResults=100,
                order="relevance",
                textFormat="plainText",
                pageToken=page_token,
            ).execute()
            _parse_page(resp2)

    except HttpError as e:
        reason = str(e)
        if "commentsDisabled" in reason or e.resp.status == 403:
            log.info(f"Comments disabled: {video_id}")
        else:
            log.warning(f"Comments error for {video_id}: {e}")

    return comments[:max_comments]


def build_transcript_api() -> None:
    """No-op — kept for API compatibility. Audio is now fetched via yt_dlp in extract_transcript_arguments."""
    return None


_AUDIO_MIME: dict[str, str] = {
    "m4a": "audio/mp4", "mp4": "audio/mp4",
    "webm": "audio/webm", "opus": "audio/ogg",
    "mp3": "audio/mpeg", "ogg": "audio/ogg",
}


def _download_audio(video_url: str, output_dir: str) -> Optional[tuple[str, str]]:
    """
    Download best audio (no ffmpeg needed) to output_dir.
    Tries multiple yt_dlp player clients sequentially.
    Returns (file_path, mime_type) or None on failure.
    """
    proxy = os.getenv("HTTPS_PROXY") or os.getenv("HTTP_PROXY") or None
    cookies_path = str(COOKIES_PATH) if COOKIES_PATH.exists() else None
    # android* clients don't support cookies; web clients use cookies
    android_clients = {"android", "android_vr", "android_embedded", "android_creator"}

    for client_hint in [["android_vr"], ["android"], ["web"], ["ios"], ["tv_embedded"], ["mweb"]]:
        ydl_opts: dict = {
            "format": "bestaudio[ext=m4a]/bestaudio[ext=webm]/bestaudio",
            "outtmpl": os.path.join(output_dir, "%(id)s.%(ext)s"),
            "quiet": True,
            "no_warnings": True,
            "extractor_args": {"youtube": {"player_client": client_hint}},
        }
        if proxy:
            ydl_opts["proxy"] = proxy
        if cookies_path and client_hint[0] not in android_clients:
            ydl_opts["cookiefile"] = cookies_path
        try:
            with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                info = ydl.extract_info(video_url, download=True)
                vid_id = info.get("id", "")
                ext = info.get("ext", "m4a")
                path = os.path.join(output_dir, f"{vid_id}.{ext}")
                if os.path.exists(path):
                    return path, _AUDIO_MIME.get(ext, "audio/mpeg")
        except Exception:
            continue
    log.warning(f"Audio download failed for all clients: {video_url}")
    return None


# ── Ranking ────────────────────────────────────────────────────────────────────

def compute_video_ranks(df: pd.DataFrame) -> pd.DataFrame:
    """
    Add video_rank column.
    video_rank = 0.5 * (likes / max_likes) + 0.5 * (views / max_views)
    Normalized within the current set of videos.
    """
    df = df.copy()
    max_likes = df["likes"].max() or 1
    max_views = df["views"].max() or 1
    df["video_rank"] = (
        0.5 * (df["likes"] / max_likes) +
        0.5 * (df["views"] / max_views)
    )
    return df


def comment_rank(video_rank: float, comment_likes: int) -> float:
    """
    comment_rank = log(1 + video_rank) * comment_likes
    log(1 + video_rank) softens the video popularity multiplier.
    """
    return math.log(1.0 + video_rank) * comment_likes


# ── Pydantic schemas for structured Gemini output ─────────────────────────────

class ArgumentItem(BaseModel):
    category: str = Field(description="One of: non_purchase, competitor, regret")
    argument: str = Field(description="Concise English argument summary (max 20 words)")
    quote: str = Field(description="Verbatim quote from audio (Korean OK, max 60 words)")
    mention_count: int = Field(description="How many times this argument theme is mentioned in the audio (1 if only once)")

class TranscriptArguments(BaseModel):
    arguments: list[ArgumentItem] = Field(default_factory=list)

class CommentItem(BaseModel):
    comment: str = Field(description="Exact verbatim comment text")
    argument: str = Field(description="Concise English argument summary (max 15 words)")
    category: str = Field(description="One of: non_purchase, competitor, regret")

class CommentArguments(BaseModel):
    items: list[CommentItem] = Field(default_factory=list)

class QuoteItem(BaseModel):
    text: str = Field(description="Original quote text (Korean or English)")
    source_url: str
    source_type: str = Field(description="transcript or comment")

class MergedArgument(BaseModel):
    argument: str = Field(description="Merged management-ready English argument title (max 20 words)")
    combined_rank: float = Field(description="Rough relative-importance estimate for sorting only; the actual displayed score is computed deterministically in Python from summed source ranks")
    quotes: list[QuoteItem] = Field(description="2-3 most illustrative quotes")
    source_count: int = Field(description="Total number of input lines (comments/transcript mentions) grouped into this merged argument, i.e. the sum of each input line's own count:N value")
    source_indices: list[int] = Field(default_factory=list, description="Every 0-based idx (from the input line list) placed in this group — not just the illustrative ones")

class MergedArguments(BaseModel):
    merged_arguments: list[MergedArgument]


# ── Gemini helpers ─────────────────────────────────────────────────────────────

def init_gemini(api_key: str) -> genai.Client:
    return genai.Client(api_key=api_key)


# Thread-safe token usage accumulator, for cost accounting across a run
# (populated by _gemini_call; read via get_token_usage()).
import threading as _threading
_token_lock = _threading.Lock()
_token_usage = {"prompt_tokens": 0, "output_tokens": 0, "calls": 0}


def get_token_usage() -> dict:
    with _token_lock:
        return dict(_token_usage)


def reset_token_usage() -> None:
    with _token_lock:
        _token_usage["prompt_tokens"] = 0
        _token_usage["output_tokens"] = 0
        _token_usage["calls"] = 0


def _gemini_call(client: genai.Client, prompt: str, schema: type, contents=None) -> Optional[object]:
    """Call Gemini with structured JSON output and exponential backoff retry."""
    payload = contents if contents is not None else prompt
    for attempt in range(GEMINI_RETRY_ATTEMPTS):
        try:
            resp = client.models.generate_content(
                model=GEMINI_MODEL,
                contents=payload,
                config=genai_types.GenerateContentConfig(
                    response_mime_type="application/json",
                    response_schema=schema,
                ),
            )
            usage = getattr(resp, "usage_metadata", None)
            if usage is not None:
                with _token_lock:
                    _token_usage["prompt_tokens"] += getattr(usage, "prompt_token_count", 0) or 0
                    _token_usage["output_tokens"] += getattr(usage, "candidates_token_count", 0) or 0
                    _token_usage["calls"] += 1
            return schema.model_validate_json(resp.text)
        except Exception as e:
            if attempt < GEMINI_RETRY_ATTEMPTS - 1:
                wait = 2 ** (attempt + 1)
                log.warning(f"Gemini error (attempt {attempt + 1}): {e} — retrying in {wait}s")
                time.sleep(wait)
            else:
                log.error(f"Gemini failed after {GEMINI_RETRY_ATTEMPTS} attempts: {e}")
    return None


# Keep parse_json for any legacy use
def parse_json(text: str) -> Optional[dict | list]:
    """Strip markdown code fences and parse JSON."""
    text = text.strip()
    text = re.sub(r"^```[a-z]*\n?", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\n?```$", "", text)
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        return None


# ── Argument extraction: transcripts ──────────────────────────────────────────

AUDIO_DOWNLOAD_DELAY = 2.0  # seconds between sequential audio downloads


def extract_transcript_arguments(
    client: genai.Client,
    video_url: str,
    title: str,
    audio_path: str,    # pre-downloaded audio file path (empty string = skip)
    car_name: str,
    mime_type: str = "audio/mp4",
) -> list[dict]:
    """
    Upload a pre-downloaded audio file to Gemini Files API and extract arguments.
    audio_path must exist. Returns [] if path is empty or upload fails.
    """
    if not audio_path or not os.path.exists(audio_path):
        return []

    video_id = video_url.split("v=")[-1].split("&")[0]

    try:
        uploaded = client.files.upload(
            file=audio_path,
            config={"mime_type": mime_type},
        )
    except Exception as e:
        log.warning(f"Gemini file upload error ({video_id}): {e}")
        return []

    try:
        prompt = (
            f"You are a Korean automotive market research analyst preparing a management briefing.\n\n"
            f"Listen to this YouTube video about the {car_name} and extract NEGATIVE arguments about the {car_name} itself.\n\n"
            f"Extract arguments in exactly these three categories — each MUST be framed negatively about "
            f"the {car_name}, never positively:\n"
            f'1. "non_purchase" – an explicit reason someone gives for NOT buying the {car_name} '
            f'(a flaw, dealbreaker, or missing feature OF THE {car_name} ITSELF)\n'
            f'2. "competitor" – an explicit statement that a NAMED alternative car is BETTER than the '
            f'{car_name} and would be chosen instead (the comparison must clearly favor the competitor, '
            f'not the {car_name})\n'
            f'3. "regret" – a negative statement from someone who ALREADY OWNS the {car_name}, expressing '
            f'regret or disappointment after purchase\n\n'
            f"Do NOT extract (skip entirely — never force these into a category):\n"
            f"- Positive statements about the {car_name} itself (e.g., praising its performance, design, "
            f"tech, or price vs rivals) — even if made while comparing to a competitor\n"
            f"- Purely negative statements about a competitor/rival brand that do not explicitly state the "
            f"{car_name} would be chosen instead\n"
            f'- Example of what to SKIP: "{car_name} has better handling and a more efficient hybrid than '
            f'Hyundai/Kia" — this is POSITIVE about the {car_name}; do not extract it under any category.\n'
            f"- General industry/market commentary not tied to a specific purchase decision\n\n"
            f"Rules:\n"
            f"- Only include arguments clearly stated or strongly implied in the audio\n"
            f"- argument: concise English sentence (max 20 words)\n"
            f"- quote: verbatim phrase from the audio (Korean OK, max 60 words)\n"
            f"- mention_count: how many times this theme appears in the audio\n"
            f"- Sort by mention_count descending within each category\n\n"
            f"Video title: {title}"
        )
        result = _gemini_call(client, prompt, TranscriptArguments, contents=[prompt, uploaded])
        if not result:
            return []
        return [
            {
                "category": arg.category,
                "argument": arg.argument,
                "quote": arg.quote,
                "mention_count": arg.mention_count,
                "source_url": video_url,
                "source_title": title,
                "source_type": "audio",
            }
            for arg in result.arguments
            if arg.category in CATEGORIES
        ]
    finally:
        try:
            client.files.delete(name=uploaded.name)
        except Exception:
            pass




# ── Argument extraction: comments ─────────────────────────────────────────────

def extract_comment_arguments(
    client: genai.Client,
    video_url: str,
    title: str,
    comments: list[dict],
    car_name: str,
) -> list[dict]:
    """
    Send comments in batches of COMMENT_BATCH_SIZE to Gemini.
    Returns list of {category, argument, comment, comment_likes, source_url, source_title, source_type}.
    """
    results: list[dict] = []

    likes_lookup: dict[str, int] = {}
    for c in comments:
        if c["comment"] not in likes_lookup:
            likes_lookup[c["comment"]] = c["likes"]

    for batch_start in range(0, len(comments), COMMENT_BATCH_SIZE):
        batch = comments[batch_start: batch_start + COMMENT_BATCH_SIZE]
        numbered = "\n".join(
            f"{i + 1}. [likes:{c['likes']}] {c['comment']}"
            for i, c in enumerate(batch)
        )
        prompt = (
            f"You are a Korean automotive market research analyst.\n\n"
            f"Below are YouTube comments from a video about the {car_name}.\n\n"
            f"Identify ONLY comments that contain arguments in these categories — each MUST be framed "
            f"negatively about the {car_name} itself, never positively:\n"
            f'1. "non_purchase" – an explicit reason NOT to buy the {car_name} (a flaw, dealbreaker, or '
            f'missing feature OF THE {car_name} ITSELF)\n'
            f'2. "competitor" – an explicit statement that a NAMED alternative car is BETTER and would be '
            f'chosen instead of the {car_name} (must clearly favor the competitor, not the {car_name})\n'
            f'3. "regret" – regret or disappointment expressed by someone who ALREADY OWNS the {car_name}\n\n'
            f"Ignore / do NOT extract:\n"
            f"- General praise, off-topic comments, questions without arguments, greetings, emojis-only\n"
            f"- Positive statements about the {car_name} itself, even when comparing it favorably to a competitor\n"
            f"- Purely negative statements about a competitor/rival brand that do not explicitly say the "
            f"{car_name} would be chosen instead\n"
            f'- Example of what to SKIP: "{car_name} has better handling and a more efficient hybrid than '
            f'Hyundai/Kia" — this is POSITIVE about the {car_name}; do not extract it under any category.\n\n'
            f"Video: {title}\nComments:\n{numbered}"
        )
        result = _gemini_call(client, prompt, CommentArguments)
        if result:
            for item in result.items:
                if item.category in CATEGORIES:
                    results.append({
                        "comment": item.comment,
                        "argument": item.argument,
                        "category": item.category,
                        "comment_likes": likes_lookup.get(item.comment, 0),
                        "source_url": video_url,
                        "source_title": title,
                        "source_type": "comment",
                    })

        time.sleep(0.3)

    return results


# ── Argument merging ───────────────────────────────────────────────────────────

MERGE_CHUNK_SIZE = 100   # args per Gemini merge call
MERGE_MAX_FINAL = 15     # hard cap on final merged arguments


def _merge_chunk(
    client: genai.Client,
    chunk: list[dict],
    category_label: str,
    car_name: str,
) -> list[dict]:
    """Merge one chunk of ≤MERGE_CHUNK_SIZE arguments via Gemini structured output."""
    def _safe_str(v) -> str:
        # Guards against NaN (float) values coming from pandas round-trips through CSV,
        # where a missing 'quote'/'comment' cell becomes float('nan') instead of "".
        return v if isinstance(v, str) else ""

    lines = [
        f"[idx:{i}|rank:{a.get('rank', 0.0):.4f}|type:{a.get('source_type','')}|url:{a.get('source_url','')}|count:{a.get('source_count', 1)}] "
        f"{a.get('argument','')} || quote: {(_safe_str(a.get('quote')) or _safe_str(a.get('comment')))[:120]}"
        for i, a in enumerate(chunk)
    ]
    prompt = (
        f"You are a senior market research analyst writing a management briefing.\n\n"
        f"Arguments from Korean YouTube videos/comments about the {car_name} — category: {category_label}.\n\n"
        f"Each line: [idx:i|rank:X|type:transcript/comment|url:...|count:N] argument || quote: ...\n"
        f"idx:i is the 0-based index of that line — use it to report exactly which lines were grouped.\n"
        f"Higher rank = more impactful source. count:N is how many individual sources that single line already represents.\n\n"
        f"Task:\n"
        f"1. Group semantically similar/overlapping arguments together\n"
        f"2. Merge each group into ONE clear English argument title (max 20 words)\n"
        f"3. Estimate combined_rank as the group's relative importance (any positive number — used only as a rough sort hint, the real score is computed separately)\n"
        f"4. Keep 2-3 most illustrative quotes per group (original Korean OK)\n"
        f"5. Set source_indices to EVERY idx value placed in that group (not just the illustrative ones)\n"
        f"6. Set source_count to the sum of count:N across every input line placed in that group\n"
        f"7. Target 5-10 merged arguments maximum — aggressively group similar themes\n"
        f"8. Sort by combined_rank descending\n\n"
        f"Arguments:\n" + "\n".join(lines)
    )
    result = _gemini_call(client, prompt, MergedArguments)
    if not result:
        return []
    n = len(chunk)
    out = []
    for m in result.merged_arguments:
        valid_idx = [i for i in m.source_indices if isinstance(i, int) and 0 <= i < n]
        members = [chunk[i] for i in valid_idx]
        # Deterministic count/rank from the actual grouped members — never trust Gemini's
        # own arithmetic. Falls back to Gemini's raw estimate only if it failed to report
        # any valid indices. combined_rank is left uncapped here; it's normalized to a
        # 0-1 scale exactly once, at the very end of merge_and_rerank, to avoid the
        # double-capping that caused many arguments to saturate at 1.0.
        source_count = sum(mm.get("source_count", 1) for mm in members) if members else max(m.source_count, 1)
        combined_rank = sum(mm.get("rank", 0.0) for mm in members) if members else max(m.combined_rank, 0.0)

        # Deterministic quote selection: pick the highest-weight (rank) members rather
        # than trusting Gemini's own quote picks, so the displayed examples always match
        # whichever sources actually drove the argument's score.
        if members:
            seen_texts: set[str] = set()
            top_quotes = []
            for mm in sorted(members, key=lambda x: x.get("rank", 0.0), reverse=True):
                text = _safe_str(mm.get("quote")) or _safe_str(mm.get("comment"))
                if not text or text in seen_texts:
                    continue
                seen_texts.add(text)
                top_quotes.append({
                    "text": text[:300],
                    "source_url": mm.get("source_url", ""),
                    "source_type": mm.get("source_type", ""),
                    "date": mm.get("date", ""),
                    "likes": mm.get("likes", ""),
                    "views": mm.get("views", ""),
                })
                if len(top_quotes) >= 3:
                    break
            quotes = top_quotes or [{"text": q.text, "source_url": q.source_url, "source_type": q.source_type} for q in m.quotes]
        else:
            quotes = [{"text": q.text, "source_url": q.source_url, "source_type": q.source_type} for q in m.quotes]

        out.append({
            "argument": m.argument,
            "combined_rank": combined_rank,
            "quotes": quotes,
            "source_count": max(source_count, 1),
            "members": members,
        })
    return out


def merge_and_rerank(
    client: genai.Client,
    arguments: list[dict],
    category: str,
    car_name: str,
) -> list[dict]:
    """
    Two-pass chunked merge:
      Pass 1 — split raw args into MERGE_CHUNK_SIZE chunks, merge each independently
      Pass 2 — merge the pass-1 results into final MERGE_MAX_FINAL arguments
    Falls back to top-N deduplication if both passes fail.
    """
    if not arguments:
        return []

    category_label = CATEGORY_LABELS[category]

    # Pass 1: chunk merge
    pass1_results: list[dict] = []
    chunks = [arguments[i:i + MERGE_CHUNK_SIZE] for i in range(0, len(arguments), MERGE_CHUNK_SIZE)]
    log.info(f"  Merge {category} — {len(arguments)} args → {len(chunks)} chunk(s)")

    for idx, chunk in enumerate(chunks):
        merged = _merge_chunk(client, chunk, category_label, car_name)
        if merged:
            for m in merged:
                pass1_results.append({
                    "argument": m["argument"],
                    "combined_rank": m["combined_rank"],
                    "quotes": m["quotes"],
                    "source_count": m["source_count"],
                    "source_items": m["members"],
                })
            log.info(f"    Chunk {idx+1}/{len(chunks)}: {len(chunk)} args → {len(merged)} merged")
        else:
            # chunk fallback: top-5 by rank
            log.warning(f"    Chunk {idx+1}/{len(chunks)}: merge failed, using top-5 fallback")
            for a in sorted(chunk, key=lambda x: x.get("rank", 0.0), reverse=True)[:5]:
                _q, _c = a.get("quote"), a.get("comment")
                quote_text = (_q if isinstance(_q, str) else "") or (_c if isinstance(_c, str) else "")
                pass1_results.append({
                    "argument": a.get("argument", ""),
                    "combined_rank": round(a.get("rank", 0.0), 4),
                    "quotes": [{"text": quote_text[:200], "source_url": a.get("source_url", ""), "source_type": a.get("source_type", "")}],
                    "source_count": a.get("source_count", 1),
                    "source_items": [a],
                })

    if not pass1_results:
        return []

    # Pass 2: final merge across chunk results (already small enough)
    log.info(f"  Merge {category} — pass 2: {len(pass1_results)} → target ≤{MERGE_MAX_FINAL}")
    if len(pass1_results) <= MERGE_MAX_FINAL:
        final = sorted(pass1_results, key=lambda x: x.get("combined_rank", 0.0), reverse=True)
    else:
        # Convert pass1 results to arg dicts compatible with _merge_chunk.
        # source_items is carried along unused by the prompt-building code, purely so
        # we can flatten it back out once pass 2 tells us which pass-1 groups merged.
        p1_as_args = [
            {
                "argument": r["argument"],
                "rank": r.get("combined_rank", 0.0),
                "quote": r["quotes"][0]["text"] if r.get("quotes") else "",
                "source_url": r["quotes"][0]["source_url"] if r.get("quotes") else "",
                "source_type": r["quotes"][0]["source_type"] if r.get("quotes") else "",
                "source_count": r.get("source_count", 1),
                "source_items": r.get("source_items", []),
            }
            for r in pass1_results
        ]
        merged2 = _merge_chunk(client, p1_as_args, category_label, car_name)
        if merged2:
            final = [
                {
                    "argument": m["argument"],
                    "combined_rank": m["combined_rank"],
                    "quotes": m["quotes"],
                    "source_count": m["source_count"],
                    "source_items": [item for member in m["members"] for item in member.get("source_items", [])],
                }
                for m in merged2
            ]
        else:
            log.warning(f"  Pass 2 merge failed for {car_name}/{category}, using pass1 top-{MERGE_MAX_FINAL}")
            final = sorted(pass1_results, key=lambda x: x.get("combined_rank", 0.0), reverse=True)

    # Re-sort by the deterministic combined_rank we just computed — Gemini's own
    # ordering of merged2 reflected its untrusted internal estimate, not the real
    # recomputed score, so results could otherwise display out of rank order.
    final = sorted(final, key=lambda x: x.get("combined_rank", 0.0), reverse=True)[:MERGE_MAX_FINAL]

    # Normalize combined_rank to 0-1 exactly once, here, relative to the top argument in
    # this category — instead of capping at each merge stage (which caused many
    # unrelated arguments to all saturate at 1.0 and become indistinguishable).
    max_rank = max((f.get("combined_rank", 0.0) for f in final), default=0.0)
    if max_rank > 0:
        for f in final:
            f["combined_rank"] = round(f.get("combined_rank", 0.0) / max_rank, 4)

    return final


# ── Docx helpers ───────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str) -> None:
    """Set background color of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_separator(doc: Document) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement("w:pageBreakBefore")
    pb.set(qn("w:val"), "0")
    pPr.append(pb)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("─" * 80)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(8)


def _add_argument_block(doc: Document, rank_idx: int, merged: dict) -> None:
    """Add one ranked argument with its quotes to the document."""
    arg_text = merged.get("argument", "N/A")
    combined_rank = merged.get("combined_rank", 0.0)
    quotes = merged.get("quotes", [])
    source_count = merged.get("source_count", 1)

    # Argument title
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run_num = p.add_run(f"{rank_idx}. ")
    run_num.bold = True
    run_num.font.size = Pt(11)
    run_main = p.add_run(arg_text)
    run_main.bold = True
    run_main.font.size = Pt(11)
    source_label = "source" if source_count == 1 else "sources"
    run_score = p.add_run(f"  [score: {combined_rank:.3f} · {source_count} {source_label}]")
    run_score.font.size = Pt(9)
    run_score.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Quotes
    for q in quotes[:3]:
        q_text = q.get("text", "").strip()
        q_url = q.get("source_url", "")
        q_type = q.get("source_type", "")
        if not q_text:
            continue

        q_para = doc.add_paragraph()
        q_para.paragraph_format.left_indent = Inches(0.4)
        q_para.paragraph_format.space_before = Pt(2)
        q_para.paragraph_format.space_after = Pt(2)

        quote_run = q_para.add_run(f'"{q_text}"')
        quote_run.italic = True
        quote_run.font.size = Pt(9.5)
        quote_run.font.color.rgb = RGBColor(0x33, 0x33, 0x55)

        meta_bits = []
        q_date = q.get("date", "")
        if q_date:
            meta_bits.append(str(q_date))
        q_views = q.get("views", "")
        if isinstance(q_views, (int, float)) and q_views:
            meta_bits.append(f"{int(q_views):,} views")
        q_likes = q.get("likes", "")
        if isinstance(q_likes, (int, float)) and q_likes:
            meta_bits.append(f"{int(q_likes):,} likes")
        meta_str = "  ·  ".join(meta_bits)

        if q_url:
            suffix = f"  [{meta_str}]" if meta_str else ""
            src_run = q_para.add_run(f"\n  → {q_url}  [{q_type}]{suffix}")
            src_run.font.size = Pt(8)
            src_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def build_docx_report(
    results_per_car: dict,
    video_counts: dict,
    comment_counts: dict,
    output_path: Path,
) -> None:
    """
    Generate the structured .docx management report.

    Structure:
      - Title + methodology
      - Per car: non_purchase + competitor reasons
      - Separate section: regret (both cars)
    """
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # ── Title ──────────────────────────────────────────────────────────────────
    title = doc.add_heading("Korean Consumer Barriers to Purchase", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Grand Koleos & Filante — YouTube Analysis")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(13)
    sub.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    date_p = doc.add_paragraph(f"Date: {pd.Timestamp.now().strftime('%B %d, %Y')}")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.runs[0].font.size = Pt(10)
    date_p.runs[0].font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    doc.add_paragraph()

    # ── Methodology ────────────────────────────────────────────────────────────
    doc.add_heading("Methodology", 1)

    total_videos = sum(video_counts.values())
    total_comments = sum(comment_counts.values())

    p = doc.add_paragraph()
    p.add_run("Data Sources: ").bold = True
    p.add_run(
        f"{total_videos} Korean YouTube videos ({total_comments} comments with ≥2 likes). "
        "Transcripts extracted via YouTube caption API (Korean preferred, English fallback). "
        "Comments filtered to top 200 per video by relevance."
    )

    p = doc.add_paragraph()
    p.add_run("Video Rank: ").bold = True
    p.add_run("0.5 × (likes / max_likes) + 0.5 × (views / max_views), normalized within each car.")

    p = doc.add_paragraph()
    p.add_run("Comment Rank: ").bold = True
    p.add_run("log(1 + video_rank) × comment_likes — softens video popularity bias while rewarding liked comments.")

    p = doc.add_paragraph()
    p.add_run("Argument Extraction: ").bold = True
    p.add_run(
        f"{GEMINI_MODEL} extracts arguments from transcripts (per video) "
        "and from comment batches of 50. Arguments are then semantically clustered "
        "and re-ranked by Gemini, with ranks summed across matching sources."
    )

    for car_key in ["koleos", "filante"]:
        p = doc.add_paragraph()
        p.add_run(f"{CAR_NAMES[car_key]}: ").bold = True
        p.add_run(f"{video_counts.get(car_key, 0)} videos, {comment_counts.get(car_key, 0)} qualifying comments.")

    doc.add_paragraph()

    # ── Per-car non-purchase + competitor sections ──────────────────────────────
    for car_key, car_name in CAR_NAMES.items():
        car_results = results_per_car.get(car_key, {})

        # Car header
        doc.add_heading(car_name, 1)

        for category in ["non_purchase", "competitor"]:
            label = CATEGORY_LABELS[category]
            merged_args = car_results.get(category, [])

            doc.add_heading(label, 2)

            if not merged_args:
                doc.add_paragraph("No significant arguments found for this category.")
            else:
                doc.add_paragraph(f"{len(merged_args)} argument(s) identified, ranked by combined score.")
                for idx, merged in enumerate(merged_args, 1):
                    _add_argument_block(doc, idx, merged)

            doc.add_paragraph()

        _add_separator(doc)

    # ── Post-Purchase Regret (separate section) ─────────────────────────────────
    doc.add_heading("Post-Purchase Regret", 1)
    doc.add_paragraph(
        "The following arguments represent reasons why consumers who already purchased "
        "the Grand Koleos or Filante express regret. These are distinct from purchase "
        "barriers and reflect real ownership pain points."
    )
    doc.add_paragraph()

    for car_key, car_name in CAR_NAMES.items():
        car_results = results_per_car.get(car_key, {})
        merged_args = car_results.get("regret", [])

        doc.add_heading(f"{car_name} — Regret Reasons", 2)

        if not merged_args:
            doc.add_paragraph("No significant regret arguments found.")
        else:
            doc.add_paragraph(f"{len(merged_args)} argument(s) identified, ranked by combined score.")
            for idx, merged in enumerate(merged_args, 1):
                _add_argument_block(doc, idx, merged)

        doc.add_paragraph()

    doc.save(str(output_path))
    log.info(f"Report saved → {output_path}")


def build_sources_csv(results_per_car: dict, output_path: Path) -> None:
    """
    Deep-dive / verification export: one row per raw source (comment or transcript
    mention) grouped into each merged argument shown in the report. Lets a reviewer
    check that a merged argument's score/source_count isn't overcounted, and trace
    any quote back to its original video/comment.
    """
    rows = []
    for car_key, car_name in CAR_NAMES.items():
        car_results = results_per_car.get(car_key, {})
        for category in CATEGORIES:
            merged_args = car_results.get(category, [])
            for rank_idx, merged in enumerate(merged_args, 1):
                arg_title = merged.get("argument", "N/A")
                combined_rank = merged.get("combined_rank", 0.0)
                declared_count = merged.get("source_count", 1)
                source_items = merged.get("source_items", [])
                if not source_items:
                    rows.append({
                        "car": car_name,
                        "category": CATEGORY_LABELS.get(category, category),
                        "argument_rank": rank_idx,
                        "argument": arg_title,
                        "combined_rank": combined_rank,
                        "declared_source_count": declared_count,
                        "actual_source_rows": 0,
                        "source_index": "",
                        "source_type": "",
                        "source_rank": "",
                        "source_date": "",
                        "source_likes": "",
                        "source_views": "",
                        "source_url": "",
                        "quote_or_comment": "",
                    })
                    continue
                for s_idx, item in enumerate(source_items, 1):
                    _q, _c = item.get("quote"), item.get("comment")
                    text = (_q if isinstance(_q, str) else "") or (_c if isinstance(_c, str) else "")
                    rows.append({
                        "car": car_name,
                        "category": CATEGORY_LABELS.get(category, category),
                        "argument_rank": rank_idx,
                        "argument": arg_title,
                        "combined_rank": combined_rank,
                        "declared_source_count": declared_count,
                        "actual_source_rows": len(source_items),
                        "source_index": s_idx,
                        "source_type": item.get("source_type", ""),
                        "source_rank": item.get("rank", ""),
                        "source_date": item.get("date", ""),
                        "source_likes": item.get("likes", ""),
                        "source_views": item.get("views", ""),
                        "source_url": item.get("source_url", ""),
                        "quote_or_comment": text,
                    })

    pd.DataFrame(rows).to_csv(output_path, index=False, encoding="utf-8-sig")
    log.info(f"Sources deep-dive CSV saved → {output_path}")


def _load_json_cache(path: Path) -> dict:
    """Load a per-car extraction cache (already-processed videos/comments), or start fresh."""
    if path.exists():
        try:
            data = json.loads(path.read_text(encoding="utf-8"))
            data.setdefault("videos", {})
            data.setdefault("comments", {})
            return data
        except Exception:
            log.warning(f"Failed to load cache {path}, starting fresh.")
    return {"videos": {}, "comments": {}}


def _save_json_cache(path: Path, data: dict) -> None:
    path.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")


def build_evolution_graphs(results_per_car: dict, output_dir: Path, top_n: int = 5) -> None:
    """
    For each car/category, plot the weekly-bucketed weighted mention volume of the
    top-N merged arguments (by combined_rank) over time. Weight per mention is the
    same 'rank' score used for report ranking (video_rank/comment_rank based), so
    highly-liked/viewed sources count for more than a single generic mention.
    """
    output_dir.mkdir(parents=True, exist_ok=True)

    for car_key, car_name in CAR_NAMES.items():
        car_results = results_per_car.get(car_key, {})
        for category in CATEGORIES:
            merged_args = car_results.get(category, [])[:top_n]
            if not merged_args:
                continue

            fig, ax = plt.subplots(figsize=(11, 6))
            has_data = False

            for merged in merged_args:
                arg_title = merged.get("argument", "N/A")
                source_items = merged.get("source_items", [])
                rows = []
                for item in source_items:
                    date_str = item.get("date", "")
                    weight = item.get("rank", 0.0)
                    if not date_str:
                        continue
                    try:
                        ts = pd.to_datetime(date_str)
                    except (ValueError, TypeError):
                        continue
                    rows.append({"date": ts, "weight": weight})

                if not rows:
                    continue

                s_df = pd.DataFrame(rows)
                s_df["week"] = s_df["date"].dt.to_period("W").dt.start_time
                weekly = s_df.groupby("week")["weight"].sum().sort_index()
                if weekly.empty:
                    continue

                label = (arg_title[:60] + "…") if len(arg_title) > 60 else arg_title
                ax.plot(weekly.index, weekly.values, marker="o", markersize=3, linewidth=1.5, label=label)
                has_data = True

            if not has_data:
                plt.close(fig)
                continue

            label_cat = CATEGORY_LABELS.get(category, category)
            ax.set_title(f"{car_name} — {label_cat}\nWeighted mention volume over time (top {top_n} arguments)")
            ax.set_xlabel("Week")
            ax.set_ylabel("Weighted mentions (summed source rank)")
            ax.legend(fontsize=7, loc="upper left", bbox_to_anchor=(1.01, 1.0))
            ax.grid(True, alpha=0.3)
            fig.autofmt_xdate()
            fig.tight_layout()

            out_path = output_dir / f"{car_key}_{category}_evolution.png"
            fig.savefig(out_path, dpi=150)
            plt.close(fig)
            log.info(f"  Graph saved → {out_path}")


# ── Main ───────────────────────────────────────────────────────────────────────

def main() -> None:
    yt = build_youtube(YOUTUBE_API_KEY)
    gemini = init_gemini(GOOGLE_API_KEY)
    reset_token_usage()

    results_per_car: dict[str, dict[str, list]] = {}
    video_counts: dict[str, int] = {}
    comment_counts: dict[str, int] = {}

    for car_key, car_name in CAR_NAMES.items():
        log.info(f"\n{'=' * 60}\nProcessing: {car_name}\n{'=' * 60}")

        cache_path = CACHE_DIR / f"{car_key}_cache.json"
        cache = _load_json_cache(cache_path)
        video_cache: dict = cache["videos"]
        comment_cache: dict = cache["comments"]

        # ── Step 1: Search videos ──────────────────────────────────────────────
        log.info("Step 1: Searching YouTube for relevant videos...")
        video_dict: dict[str, str] = {}
        for query in SEARCH_QUERIES[car_key]:
            found = search_videos(yt, query)
            new = {url: vid for url, vid in found.items() if url not in video_dict}
            video_dict.update(new)
            log.info(f"  '{query}' → {len(found)} videos ({len(new)} new)")
            time.sleep(0.2)

        log.info(f"Total unique videos found: {len(video_dict)}")

        # ── Step 2: Fetch video details ────────────────────────────────────────
        log.info("Step 2: Fetching video details...")
        video_details: list[dict] = []
        for url, video_id in video_dict.items():
            details = get_video_details(yt, video_id)
            if details:
                video_details.append(details)
            time.sleep(0.05)

        if not video_details:
            log.warning(f"No video details returned for {car_name} — skipping.")
            results_per_car[car_key] = {c: [] for c in CATEGORIES}
            video_counts[car_key] = 0
            comment_counts[car_key] = 0
            continue

        videos_df = pd.DataFrame(video_details)
        videos_df = compute_video_ranks(videos_df)
        video_counts[car_key] = len(videos_df)

        # Save video metadata
        videos_df.to_csv(OUTPUT_DIR / f"{car_key}_videos.csv", index=False, encoding="utf-8-sig")
        log.info(f"Videos with details: {len(videos_df)}")

        # ── Step 3: Fetch comments ─────────────────────────────────────────────
        log.info("Step 3: Fetching comments...")
        video_data: list[dict] = []
        total_qualifying_comments = 0

        for _, row in videos_df.iterrows():
            vid_id = row["video_id"]
            v_rank = float(row["video_rank"])

            raw_comments = fetch_comments(yt, vid_id)

            # Filter + rank comments
            qualifying = [
                {**c, "rank": comment_rank(v_rank, c["likes"]), "video_rank": v_rank}
                for c in raw_comments
                if c["likes"] >= MIN_COMMENT_LIKES
            ]
            total_qualifying_comments += len(qualifying)

            video_data.append({
                "video_id": vid_id,
                "url": row["url"],
                "title": row["title"],
                "video_rank": v_rank,
                "published_at": row["published_at"],
                "views": int(row["views"]),
                "likes": int(row["likes"]),
                "comments": qualifying,
            })

            log.info(f"  [{len(qualifying)} comments] {row['title'][:65]}")
            time.sleep(0.2)

        comment_counts[car_key] = total_qualifying_comments
        log.info(f"Comments qualifying (≥{MIN_COMMENT_LIKES} likes): {total_qualifying_comments}")

        # Save comments to CSV immediately (crash safety)
        comment_rows = [
            {**c, "video_title": vd["title"]}
            for vd in video_data for c in vd["comments"]
        ]
        if comment_rows:
            pd.DataFrame(comment_rows).to_csv(
                OUTPUT_DIR / f"{car_key}_comments.csv", index=False, encoding="utf-8-sig"
            )
        log.info(f"Comments saved to CSV.")

        # ── Step 3b: split comments into already-processed (cached) vs new ─────
        # Videos whose audio was already extracted in a previous run are skipped
        # entirely for transcript extraction below, but we still fetch + diff
        # comments every run since new comments can appear under old videos.
        # comment_cache is keyed by YouTube commentThread id, so re-fetched
        # duplicates across runs are naturally never double-counted.
        for vd in video_data:
            vd["new_comments"] = [c for c in vd["comments"] if c["comment_id"] not in comment_cache]

        n_new_comments = sum(len(vd["new_comments"]) for vd in video_data)
        n_cached_comments = total_qualifying_comments - n_new_comments
        log.info(f"  Comments: {n_new_comments} new (→ Gemini), {n_cached_comments} already processed (cached)")

        # ── Step 4: Extract arguments in parallel ──────────────────────────────
        videos_needing_audio = [vd for vd in video_data if vd["video_id"] not in video_cache]
        videos_cached_audio = len(video_data) - len(videos_needing_audio)
        log.info(
            f"Step 4: Extracting arguments with Gemini ({GEMINI_WORKERS} workers)...\n"
            f"  Audio jobs: {len(videos_needing_audio)} new, {videos_cached_audio} cached (skipped) | "
            f"Comment-batch jobs: {sum(1 for vd in video_data if vd['new_comments'])}"
        )

        # Phase 4a: sequential audio downloads (avoids hammering YouTube) — only for new videos
        log.info("  Phase 4a: downloading audio sequentially...")
        audio_dir = OUTPUT_DIR / f"{car_key}_audio"
        audio_dir.mkdir(exist_ok=True)
        for i, vd in enumerate(videos_needing_audio):
            result = _download_audio(vd["url"], str(audio_dir))
            if result:
                vd["audio_path"], vd["audio_mime"] = result
                log.info(f"  [{i+1}/{len(videos_needing_audio)} ✓] {vd['title'][:65]}")
            else:
                vd["audio_path"], vd["audio_mime"] = "", "audio/mp4"
                log.info(f"  [{i+1}/{len(videos_needing_audio)} ✗] {vd['title'][:65]}")
            time.sleep(AUDIO_DOWNLOAD_DELAY)

        raw_arguments: dict[str, list[dict]] = {c: [] for c in CATEGORIES}

        def _attach_transcript_meta(arg: dict, vd: dict) -> dict:
            mc = arg.pop("mention_count", 1)
            arg["rank"] = vd["video_rank"] * math.log(1 + mc)
            arg["mention_count"] = mc
            arg["date"] = vd["published_at"]
            arg["likes"] = vd["likes"]
            arg["views"] = vd["views"]
            return arg

        def _attach_comment_meta(arg: dict, meta: dict) -> dict:
            arg["rank"] = meta["rank"]
            arg["date"] = meta["published_at"]
            arg["likes"] = meta["likes"]
            arg["views"] = ""
            return arg

        # Reuse cached transcript arguments for videos already processed in a past run
        for vd in video_data:
            if vd["video_id"] in video_cache:
                for cached_arg in video_cache[vd["video_id"]].get("transcript_args", []):
                    arg = _attach_transcript_meta(dict(cached_arg), vd)
                    cat = arg.get("category")
                    if cat in raw_arguments:
                        raw_arguments[cat].append(arg)

        # Reuse cached comment arguments for comments already processed in a past run
        for vd in video_data:
            for c in vd["comments"]:
                cached = comment_cache.get(c["comment_id"])
                if cached is None:
                    continue
                cached_arg = cached.get("argument")
                if not cached_arg:
                    continue
                arg = dict(cached_arg)
                arg["comment"] = c["comment"]
                arg["source_url"] = vd["url"]
                arg["source_title"] = vd["title"]
                arg["source_type"] = "comment"
                arg = _attach_comment_meta(arg, c)
                cat = arg.get("category")
                if cat in raw_arguments:
                    raw_arguments[cat].append(arg)

        audio_futures: dict = {}
        comment_futures: dict = {}

        import threading
        _lock = threading.Lock()
        t_done = 0
        c_done = 0
        n_comment_jobs = sum(1 for vd in video_data if vd["new_comments"])

        # Phase 4b: parallel Gemini extraction (audio upload + comment batches) — new work only
        log.info("  Phase 4b: parallel Gemini extraction...")
        with ThreadPoolExecutor(max_workers=GEMINI_WORKERS) as executor:
            for vd in videos_needing_audio:
                if vd.get("audio_path"):
                    f = executor.submit(
                        extract_transcript_arguments,
                        gemini, vd["url"], vd["title"], vd["audio_path"], car_name, vd["audio_mime"],
                    )
                    audio_futures[f] = vd

            for vd in video_data:
                if vd["new_comments"]:
                    f = executor.submit(
                        extract_comment_arguments,
                        gemini, vd["url"], vd["title"], vd["new_comments"], car_name,
                    )
                    comment_futures[f] = vd

            for future in as_completed(audio_futures):
                vd = audio_futures[future]
                try:
                    args = future.result()
                    # Cache the raw (pre-metadata) extraction so a future run can skip
                    # this video's audio download + Gemini call entirely.
                    video_cache[vd["video_id"]] = {"transcript_args": [dict(a) for a in args]}
                    found = {cat: 0 for cat in CATEGORIES}
                    for arg in args:
                        arg = _attach_transcript_meta(arg, vd)
                        cat = arg.get("category")
                        if cat in raw_arguments:
                            with _lock:
                                raw_arguments[cat].append(arg)
                            found[cat] += 1
                    with _lock:
                        t_done += 1
                        done_now = t_done
                    summary = ", ".join(f"{cat}:{n}" for cat, n in found.items() if n)
                    log.info(
                        f"  [audio {done_now}/{len(videos_needing_audio)}] "
                        f"{vd['title'][:55]} → {summary or 'no args'}"
                    )
                except Exception as e:
                    log.warning(f"Audio future error ({vd['url']}): {e}")

            for future in as_completed(comment_futures):
                vd = comment_futures[future]
                try:
                    args = future.result()
                    meta_lookup = {c["comment"]: c for c in vd["new_comments"]}
                    matched_ids: set[str] = set()
                    found = {cat: 0 for cat in CATEGORIES}
                    for arg in args:
                        meta = meta_lookup.get(arg.get("comment", ""))
                        if not meta:
                            continue
                        matched_ids.add(meta["comment_id"])
                        arg = _attach_comment_meta(arg, meta)
                        cat = arg.get("category")
                        if cat in raw_arguments:
                            with _lock:
                                raw_arguments[cat].append(arg)
                            found[cat] += 1
                        comment_cache[meta["comment_id"]] = {
                            "argument": {"category": arg.get("category"), "argument": arg.get("argument")},
                        }
                    # Mark every new comment that produced no argument as processed too,
                    # so it isn't resent to Gemini on the next run.
                    for c in vd["new_comments"]:
                        if c["comment_id"] not in matched_ids:
                            comment_cache[c["comment_id"]] = {"argument": None}
                    with _lock:
                        c_done += 1
                        done_now = c_done
                    n_batches = math.ceil(len(vd["new_comments"]) / COMMENT_BATCH_SIZE)
                    summary = ", ".join(f"{cat}:{n}" for cat, n in found.items() if n)
                    log.info(
                        f"  [comments  {done_now}/{n_comment_jobs}] "
                        f"{vd['title'][:55]} ({n_batches} batch{'es' if n_batches>1 else ''}) "
                        f"→ {summary or 'no args'}"
                    )
                except Exception as e:
                    log.warning(f"Comment future error ({vd['url']}): {e}")

        # Persist cache updates (audio + comments) for this car right away.
        _save_json_cache(cache_path, cache)

        for cat, args in raw_arguments.items():
            log.info(f"  Raw arguments — {cat}: {len(args)}")

        # Save raw arguments to CSV
        all_raw = []
        for cat, args in raw_arguments.items():
            for a in args:
                all_raw.append({**a, "category": cat})
        if all_raw:
            pd.DataFrame(all_raw).to_csv(
                OUTPUT_DIR / f"{car_key}_raw_arguments.csv", index=False, encoding="utf-8-sig"
            )

        # ── Step 5: Merge and re-rank ──────────────────────────────────────────
        log.info("Step 5: Merging and re-ranking arguments with Gemini...")
        merged_results: dict[str, list] = {}

        with ThreadPoolExecutor(max_workers=3) as executor:
            merge_futures = {
                executor.submit(merge_and_rerank, gemini, raw_arguments[cat], cat, car_name): cat
                for cat in CATEGORIES
                if raw_arguments[cat]
            }
            for f in as_completed(merge_futures):
                cat = merge_futures[f]
                try:
                    merged = f.result()
                    merged_results[cat] = merged
                    log.info(f"  Merged — {cat}: {len(merged)} arguments")
                except Exception as e:
                    log.warning(f"Merge future error ({cat}): {e}")
                    merged_results[cat] = []

        # Fill empty categories
        for cat in CATEGORIES:
            if cat not in merged_results:
                merged_results[cat] = []

        results_per_car[car_key] = merged_results

    # ── Step 6: Generate .docx report ─────────────────────────────────────────
    log.info("\nStep 6: Generating .docx report...")
    output_path = OUTPUT_DIR / "koleos_filante_non_purchase_analysis.docx"
    build_docx_report(results_per_car, video_counts, comment_counts, output_path)

    sources_path = OUTPUT_DIR / "koleos_filante_sources_deep_dive.csv"
    build_sources_csv(results_per_car, sources_path)

    log.info("Step 7: Generating evolution graphs...")
    build_evolution_graphs(results_per_car, OUTPUT_DIR / "graphs")

    usage = get_token_usage()
    cost = (
        usage["prompt_tokens"] / 1_000_000 * INPUT_PRICE_PER_1M
        + usage["output_tokens"] / 1_000_000 * OUTPUT_PRICE_PER_1M
    )
    log.info(
        f"TOKEN USAGE: calls={usage['calls']} prompt_tokens={usage['prompt_tokens']:,} "
        f"output_tokens={usage['output_tokens']:,} total_tokens={usage['prompt_tokens'] + usage['output_tokens']:,}"
    )
    log.info(f"ESTIMATED COST (gemini-3.1-flash-lite standard rate): ${cost:.4f}")

    log.info(f"\nDone. Report: {output_path}")


if __name__ == "__main__":
    main()
