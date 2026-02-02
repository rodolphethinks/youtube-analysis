"""
Report generation module for creating analysis outputs.
Generates Word documents, Excel files, and summary reports.
"""

import json
from typing import Dict, List, Optional, Any, TYPE_CHECKING
from pathlib import Path
from datetime import datetime

import pandas as pd

if TYPE_CHECKING:
    from .config import CarModel, PipelineConfig
    from .analysis import VideoAnalysis, CommentAnalysis

from .analysis import GeminiClient, analysis_to_dataframe


class ReportGenerator:
    """Generate comprehensive analysis reports."""
    
    def __init__(
        self, 
        gemini_client: GeminiClient, 
        output_dir: str = "output"
    ):
        self.client = gemini_client
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
    
    def generate_summary_report(
        self, 
        analysis_df: pd.DataFrame,
        car_model: "CarModel"
    ) -> str:
        """Generate a comprehensive summary report using AI."""
        columns = [
            'Overall Sentiment', 'Key Strengths', 'Key Weaknesses',
            'Competitor Mentions', 'Comparison Summary', 'Trends',
            'Battery Performance', 'Noise Levels', 'Competitor Perception',
            'Chinese Brand Mentions', 'Final Verdict'
        ]
        
        extracted_data = {
            col: analysis_df[col].dropna().tolist() 
            for col in columns if col in analysis_df.columns
        }
        
        prompt = f'''You are an expert automotive market analyst. Given the following data extracted from YouTube video reviews about the {car_model.company} {car_model.model}, provide a comprehensive analysis report.

Identify the most common themes, significant insights, and key takeaways for each aspect.

Data:
{json.dumps(extracted_data, indent=2, ensure_ascii=False)}

Provide the analysis in the following structured format:

## Executive Summary
[2-3 sentence overview of the overall reception]

## Overall Sentiment Analysis
- Distribution of positive/neutral/negative reviews
- Key drivers of sentiment

## Key Strengths
- Top mentioned positive aspects with frequency
- Analysis of what resonates with reviewers

## Key Weaknesses
- Top mentioned concerns with frequency
- Areas for improvement

## Competitive Landscape
- Main competitors mentioned
- How the {car_model.model} compares

## Market Trends
- Emerging themes in the reviews
- Consumer preferences indicated

## Technical Aspects
- Battery performance feedback
- Noise level observations
- Other technical notes

## Recommendations
- For potential buyers
- For the manufacturer

## Conclusion
[Final summary and outlook]
'''
        
        response = self.client.generate(prompt)
        return response
    
    def save_to_word(
        self, 
        report_text: str, 
        filename: str,
        car_model: "CarModel"
    ) -> Path:
        """Save report to Word document with proper markdown parsing."""
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
        except ImportError:
            print("python-docx not installed. Saving as text file instead.")
            return self.save_to_text(report_text, filename.replace('.docx', '.txt'))
        
        doc = Document()
        
        # Set up styles
        styles = doc.styles
        
        # Title
        title = doc.add_heading(
            f'{car_model.company} {car_model.model} - Video Analysis Report', 
            level=0
        )
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacer
        
        # Parse markdown and convert to docx
        self._parse_markdown_to_docx(doc, report_text)
        
        output_path = self.output_dir / filename
        doc.save(str(output_path))
        print(f"Report saved to: {output_path}")
        return output_path
    
    def _parse_markdown_to_docx(self, doc, markdown_text: str):
        """Parse markdown text and add formatted content to document."""
        import re
        from docx.shared import Pt, RGBColor
        
        lines = markdown_text.split('\n')
        i = 0
        in_list = False
        list_indent_level = 0
        
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            
            # Skip empty lines but end lists
            if not stripped:
                in_list = False
                i += 1
                continue
            
            # Headers
            if stripped.startswith('#### '):
                doc.add_heading(stripped[5:], level=4)
                in_list = False
            elif stripped.startswith('### '):
                doc.add_heading(stripped[4:], level=3)
                in_list = False
            elif stripped.startswith('## '):
                doc.add_heading(stripped[3:], level=2)
                in_list = False
            elif stripped.startswith('# '):
                doc.add_heading(stripped[2:], level=1)
                in_list = False
            
            # Bullet points (-, *, or numbered)
            elif re.match(r'^[-*]\s+', stripped) or re.match(r'^\d+\.\s+', stripped):
                # Determine indent level by leading spaces
                leading_spaces = len(line) - len(line.lstrip())
                indent_level = leading_spaces // 2
                
                # Extract bullet content
                if re.match(r'^[-*]\s+', stripped):
                    content = re.sub(r'^[-*]\s+', '', stripped)
                    style = 'List Bullet'
                    if indent_level > 0:
                        style = 'List Bullet 2' if indent_level == 1 else 'List Bullet 3'
                else:
                    content = re.sub(r'^\d+\.\s+', '', stripped)
                    style = 'List Number'
                    if indent_level > 0:
                        style = 'List Number 2' if indent_level == 1 else 'List Number 3'
                
                # Add bullet with formatting
                try:
                    para = doc.add_paragraph(style=style)
                except:
                    para = doc.add_paragraph(style='List Bullet')
                self._add_formatted_text(para, content)
                in_list = True
            
            # Horizontal rule
            elif stripped in ['---', '***', '___']:
                doc.add_paragraph('─' * 50)
            
            # Regular paragraph
            else:
                para = doc.add_paragraph()
                self._add_formatted_text(para, stripped)
                in_list = False
            
            i += 1
    
    def _add_formatted_text(self, paragraph, text: str):
        """Add text to paragraph with bold, italic, and other markdown formatting."""
        import re
        from docx.shared import Pt, RGBColor
        
        # Pattern to find **bold**, *italic*, `code`, and ***bold italic***
        pattern = r'(\*\*\*[^*]+\*\*\*|\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)'
        
        parts = re.split(pattern, text)
        
        for part in parts:
            if not part:
                continue
            
            # Bold italic (***text***)
            if part.startswith('***') and part.endswith('***'):
                run = paragraph.add_run(part[3:-3])
                run.bold = True
                run.italic = True
            # Bold (**text**)
            elif part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            # Italic (*text*)
            elif part.startswith('*') and part.endswith('*'):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            # Code (`text`)
            elif part.startswith('`') and part.endswith('`'):
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x64, 0x64, 0x64)
            # Regular text
            else:
                paragraph.add_run(part)
    
    def save_to_text(self, report_text: str, filename: str) -> Path:
        """Save report to text file."""
        output_path = self.output_dir / filename
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(report_text)
        print(f"Report saved to: {output_path}")
        return output_path
    
    def save_to_excel(
        self,
        videos_df: pd.DataFrame,
        analysis_df: pd.DataFrame,
        comments_df: Optional[pd.DataFrame] = None,
        filename: str = "analysis.xlsx",
        car_model: Optional["CarModel"] = None
    ) -> Path:
        """Save analysis data to Excel with multiple sheets."""
        output_path = self.output_dir / filename
        
        with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
            # Merge video details with analysis
            if 'Video URL' in videos_df.columns and 'Video URL' in analysis_df.columns:
                merged_df = analysis_df.merge(
                    videos_df, 
                    on='Video URL', 
                    how='inner',
                    suffixes=('', '_video')
                )
                
                # Clean up title for Excel
                if 'Title' in merged_df.columns:
                    merged_df['Title'] = merged_df['Title'].str.replace('"', "'", regex=False)
                
                # Sort by views
                if 'Views' in merged_df.columns:
                    merged_df = merged_df.sort_values('Views', ascending=False)
                
                merged_df.to_excel(writer, sheet_name='Analysis', index=False)
            else:
                analysis_df.to_excel(writer, sheet_name='Analysis', index=False)
            
            # Video details sheet
            videos_df.to_excel(writer, sheet_name='Videos', index=False)
            
            # Comments sheet (if available)
            if comments_df is not None and not comments_df.empty:
                comments_df.to_excel(writer, sheet_name='Comments', index=False)
        
        print(f"Excel file saved to: {output_path}")
        return output_path
    
    def save_comments_csv(
        self,
        comments_df: pd.DataFrame,
        filename: str = "comments.csv"
    ) -> Path:
        """Save comments to CSV."""
        output_path = self.output_dir / filename
        comments_df.to_csv(output_path, index=False, encoding='utf-8')
        print(f"Comments saved to: {output_path}")
        return output_path


class MultiModelReportGenerator:
    """Generate comparison reports for multiple car models."""
    
    def __init__(self, gemini_client: GeminiClient, output_dir: str = "output"):
        self.client = gemini_client
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
    
    def generate_comparison_excel(
        self,
        model_analyses: Dict[str, pd.DataFrame],
        filename: str = "comparison.xlsx"
    ) -> Path:
        """Generate comparison Excel with multiple sheets."""
        output_path = self.output_dir / filename
        
        with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
            for model_name, df in model_analyses.items():
                sheet_name = model_name[:31]  # Excel sheet name limit
                df.to_excel(writer, sheet_name=sheet_name, index=False)
        
        print(f"Comparison saved to: {output_path}")
        return output_path
    
    def generate_sentiment_comparison(
        self,
        model_analyses: Dict[str, pd.DataFrame]
    ) -> pd.DataFrame:
        """Generate sentiment comparison across models."""
        sentiment_data = []
        
        for model_name, df in model_analyses.items():
            if 'Overall Sentiment' in df.columns:
                counts = df['Overall Sentiment'].value_counts()
                sentiment_data.append({
                    'Model': model_name,
                    'Positive': counts.get('Positive', 0),
                    'Neutral': counts.get('Neutral', 0),
                    'Negative': counts.get('Negative', 0),
                    'Total': len(df)
                })
        
        return pd.DataFrame(sentiment_data)
    
    def visualize_sentiment(
        self,
        sentiment_df: pd.DataFrame,
        filename: str = "sentiment_comparison.png"
    ) -> Optional[Path]:
        """Create sentiment distribution visualization."""
        try:
            import matplotlib.pyplot as plt
            import seaborn as sns
        except ImportError:
            print("matplotlib/seaborn not installed. Skipping visualization.")
            return None
        
        plt.figure(figsize=(12, 8))
        
        sentiment_df_plot = sentiment_df.set_index('Model')[['Positive', 'Neutral', 'Negative']]
        colors = sns.color_palette("coolwarm", 3)
        
        sentiment_df_plot.plot(kind='bar', width=0.8, color=colors, edgecolor='black')
        
        plt.title('Overall Sentiment Distribution Comparison', fontsize=16, fontweight='bold')
        plt.xlabel('Car Model', fontsize=14)
        plt.ylabel('Number of Reviews', fontsize=12)
        plt.xticks(rotation=45, ha='right', fontsize=12)
        plt.grid(axis='y', linestyle='--', alpha=0.7)
        plt.legend(title='Sentiment', fontsize=10, title_fontsize=12)
        plt.tight_layout()
        
        output_path = self.output_dir / filename
        plt.savefig(output_path, dpi=300, bbox_inches='tight')
        plt.close()
        
        print(f"Visualization saved to: {output_path}")
        return output_path
